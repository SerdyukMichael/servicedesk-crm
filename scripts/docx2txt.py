#!/usr/bin/env python3
"""
Конвертер .docx → plain text для git diff.
Использование: python docx2txt.py <file.docx>
"""
import sys
from docx import Document


def docx_to_text(path: str) -> str:
    doc = Document(path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ""

        # Заголовки — выделяем маркером для наглядности в diff
        if style.startswith("Heading 1"):
            lines.append(f"\n{'='*60}")
            lines.append(text)
            lines.append('='*60)
        elif style.startswith("Heading 2"):
            lines.append(f"\n{'─'*40}")
            lines.append(text)
            lines.append('─'*40)
        elif style.startswith("Heading"):
            lines.append(f"\n### {text}")
        elif text:
            lines.append(text)

    # Таблицы
    for table in doc.tables:
        lines.append("")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append(" | ".join(cells))
        lines.append("")

    return "\n".join(lines)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: docx2txt.py <file.docx>", file=sys.stderr)
        sys.exit(1)

    try:
        # UTF-8 stdout — важно для Windows, где cp1251 по умолчанию
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        text = docx_to_text(sys.argv[1])
        try:
            sys.stdout.write(text + "\n")
            sys.stdout.flush()
        except OSError:
            pass  # git перехватывает поток — ошибка сброса буфера безопасна
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
